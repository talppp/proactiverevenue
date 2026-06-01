"""
Urgency Taxonomy + Topic Taxonomy + Team Tagging document.
Tenant/Buyer/Seller use cases, with concrete examples per row and team-routing
defaults. Emits DOCX + JSON drop-in for the classifier prompts.
"""
import os, json
from docx import Document
from docx.shared import Inches, Pt, RGBColor

OUT = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router\06_Taxonomy"
DOCX = os.path.join(OUT, "Urgency_Topic_Taxonomy_and_Team_Tagging.docx")
JSON = os.path.join(OUT, "taxonomy.json")

# ---------------------------------------------------------------------------
# Taxonomy data
# ---------------------------------------------------------------------------
URGENCY = [
    {"code":"P0","name":"Emergency","sla":"15 min","example":"Tenant calling from a flooding apartment; lockout at 11pm; gas leak smell.","action":"Push + SMS to Nathalie immediately; ring after 5 min unread."},
    {"code":"P1","name":"Same-day","sla":"4 hours","example":"Buyer offer with same-day deadline; closing tomorrow needs a doc.","action":"Push + queue card to assigned owner."},
    {"code":"P2","name":"24-hour","sla":"24 hours","example":"Showing reschedule for next week; standard contract question.","action":"Queue card + appears in next daily digest."},
    {"code":"P3","name":"Backlog","sla":"72 hours","example":"FYI updates; future inquiries with 'no rush'.","action":"Batched digest only."},
]

TOPICS = [
  # (topic_code, label, primary_owner, fallback, examples)
  ("buyer_lead",    "Buyer lead (new prospect)",        "Nathalie", "Fabi (non-core county)",
   "First contact buyer: 'Are you the agent for 1234 Peach Ln?' / 'Looking for 3BR in Sandy Springs.'"),
  ("seller_lead",   "Seller lead (new listing)",         "Nathalie", "Fabi (non-core county)",
   "'I want to sell my house in Marietta.' / 'How much could I get for...'"),
  ("showing",       "Showing logistics",                 "Nathalie", "Fabi",
   "'Can we move tomorrow's 4pm showing?' / 'I'll be 20 min late.'"),
  ("tenant_issue",  "Tenant living-condition issue",     "Nathalie", "Noa",
   "'AC not cooling' / 'water heater leaking' / 'neighbor noise complaint'."),
  ("maintenance",   "Maintenance / repair coordination", "Noa",      "Fabi",
   "Vendor scheduling, plumber dispatch, follow-up on a work order."),
  ("contract",      "Contract review/exec",              "Noa",      "Nathalie",
   "'Please review the redline' / 'Need your signature on...'"),
  ("document",      "Document request / upload",         "Noa",      "Nathalie",
   "Buyer asks for the seller's disclosure; lender requests proof of insurance."),
  ("vendor",        "Vendor outreach / quotes",          "Noa",      "Nathalie",
   "Plumber availability, landscaper invoice question, painter quote follow-up."),
  ("admin",         "General back-office admin",         "Noa",      "Nathalie",
   "Office supplies, software access, expense reimbursements."),
  ("scheduling",    "Calendar / meeting requests",       "Noa",      "Nathalie",
   "'Can we meet Thursday?' / lender wants a 3-way call."),
  ("billing",       "Invoicing / payments / deposits",   "Noa",      "Nathalie",
   "Earnest money status, vendor invoice approval, refund question."),
  ("deal",          "Active deal updates",               "Nathalie", "Noa",
   "Counter-offers, inspection negotiation, appraisal results."),
  ("vip_client",    "VIP client (Google Contacts label)","Nathalie", "Nathalie",
   "Anyone tagged VIP - hand to Nathalie regardless of topic."),
  ("personal",      "Personal / non-business",           "ARCHIVE",  "-",
   "Family chat, friend invite, doctor appointment confirmation."),
  ("marketing_promo","Marketing / promotional",          "ARCHIVE",  "-",
   "Vendor cold pitch, MailChimp campaign, generic newsletter."),
  ("spam",          "Spam / phishing",                    "ARCHIVE",  "-",
   "'Click here to claim your prize' / impersonation attempts."),
]

USE_CASES = [
    ("Real-estate BUYER", [
        ("First contact - listing inquiry",          "P2", "buyer_lead",  "Nathalie",
         "Reply with disclosure brochure + book a 15-min call."),
        ("Buyer financing question",                  "P1", "buyer_lead",  "Nathalie",
         "Loop in preferred lender; draft a reply."),
        ("Offer expires today",                        "P0", "deal",        "Nathalie",
         "Real-time push; offer template in CRM card."),
        ("Buyer asks to view a non-Apteker listing",  "P2", "buyer_lead",  "Nathalie",
         "Confirm property and county; reroute to Fabi if non-core county."),
        ("Buyer wants disclosures / seller's package", "P2", "document",    "Noa",
         "Pull from shared drive; reply via Nathalie."),
        ("Inspection results review",                  "P1", "deal",        "Nathalie",
         "Summarize defects; suggest negotiation points."),
    ]),
    ("Real-estate SELLER", [
        ("New listing inquiry from would-be seller", "P2", "seller_lead", "Nathalie",
         "Reply with CMA offer + intro call."),
        ("Photo / staging coordination",              "P2", "scheduling",  "Noa",
         "Calendar block + vendor email."),
        ("Price drop discussion",                      "P1", "deal",        "Nathalie",
         "Pull comp data; draft talking points."),
        ("Closing date conflict",                      "P1", "deal",        "Nathalie",
         "Lender + title coordination."),
        ("Seller wants to cancel listing",             "P1", "contract",    "Noa",
         "Reference cancellation clause; flag for Nathalie."),
    ]),
    ("Real-estate TENANT (existing managed property)", [
        ("AC / heat not working",                      "P0", "tenant_issue","Nathalie",
         "Dispatch vendor; escalate if temp extreme."),
        ("Leak / water damage",                         "P0", "tenant_issue","Nathalie",
         "Dispatch plumber; document with photo request."),
        ("Lockout / lost key",                          "P0", "tenant_issue","Nathalie",
         "Locksmith or emergency code."),
        ("Rent payment confusion",                      "P1", "billing",      "Noa",
         "Reconcile ACH; reply with statement."),
        ("Lease renewal question",                      "P2", "contract",    "Noa",
         "Send renewal template with rate."),
        ("Noise / neighbor complaint",                  "P2", "tenant_issue","Nathalie",
         "Document; if recurring, notify HOA."),
        ("Move-out notice",                             "P2", "contract",    "Noa",
         "Trigger move-out checklist."),
        ("Maintenance vendor follow-up",                "P2", "maintenance", "Noa",
         "Status check + ETA update to tenant."),
        ("Pest control request",                        "P2", "maintenance", "Noa",
         "Vendor dispatch + tenant communication."),
    ]),
    ("PROPERTY MANAGEMENT (Nathalie's existing managed portfolio)", [
        ("Vendor invoice approval",                    "P3", "billing",     "Noa",
         "Reconcile against work order."),
        ("HOA notice",                                  "P2", "admin",        "Noa",
         "File + notify Nathalie if action needed."),
        ("Insurance renewal",                           "P2", "admin",        "Noa",
         "Track expiry; renew."),
        ("Tax / mill levy notice",                      "P2", "admin",        "Noa",
         "File; flag if discrepancy."),
    ]),
]

TEAM_TAGS = [
    ("Nathalie","tag:nathalie","Owner; buyer/seller leads, deals, escalations, VIP handling.",
     "Core counties: Fulton, DeKalb.","After-hours P0 OK."),
    ("Fabi","tag:fabi","Agent for properties OUTSIDE core counties (Cobb, Gwinnett, Forsyth, Cherokee).",
     "Buyer/seller leads, showings, tenant issues in her counties.","After-hours OFF (P0 reroutes to Nathalie)."),
    ("Noa","tag:noa","Back-office; contracts, documents, vendor coordination, scheduling, admin, billing.",
     "County-agnostic.","After-hours OFF (queued to next morning)."),
]

# ---------------------------------------------------------------------------
# JSON dump (for classifier prompt)
# ---------------------------------------------------------------------------
out = {
    "urgency": URGENCY,
    "topics":  [{"code":c,"label":l,"primary":p,"fallback":f,"examples":e}
                for c,l,p,f,e in TOPICS],
    "use_cases": [{"persona":per, "scenarios":[
        {"scenario":s,"urgency":u,"topic":t,"owner":o,"action":a} for s,u,t,o,a in scen
    ]} for per,scen in USE_CASES],
    "team_tags":[{"name":n,"tag":t,"role":r,"counties":c,"after_hours":a}
                 for n,t,r,c,a in TEAM_TAGS],
}
with open(JSON,"w",encoding="utf-8") as f:
    json.dump(out, f, indent=2, ensure_ascii=False)

# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
doc = Document()
for s in doc.sections:
    s.left_margin = Inches(0.6); s.right_margin = Inches(0.6)
    s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)

def h1(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(18)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h2(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(13)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h3(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(11)
def body(t):
    p=doc.add_paragraph(t)
    for r in p.runs: r.font.size=Pt(10)
def bullet(t):
    p=doc.add_paragraph(t, style="List Bullet")
    for r in p.runs: r.font.size=Pt(10)
def table(rows, widths=None):
    t=doc.add_table(rows=0,cols=len(rows[0])); t.style="Light Grid Accent 1"
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row):
            cells[j].text=str(val)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size=Pt(9)
                    if i==0: r.bold=True
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Inches(w)

h1("Urgency Taxonomy, Topic Taxonomy & Team Tagging")
body("Apteker Realty  /  AI Issue Router  /  Phase I  /  Rev 1.0  /  2026-05-17")

h2("1. Urgency taxonomy")
rows=[["Code","Name","SLA","Example trigger","System action"]]
for u in URGENCY:
    rows.append([u["code"], u["name"], u["sla"], u["example"], u["action"]])
table(rows, widths=[0.4, 0.9, 0.7, 2.8, 2.5])

h2("2. Topic taxonomy")
rows=[["Code","Label","Primary owner","Fallback","Example messages"]]
for c,l,p,f,e in TOPICS:
    rows.append([c,l,p,f,e])
table(rows, widths=[1.1, 1.6, 1.0, 1.4, 3.0])

h2("3. Use cases (Tenant / Buyer / Seller / Property Mgmt)")
for persona, scenarios in USE_CASES:
    h3(persona)
    rows=[["Scenario","Urgency","Topic","Owner","Next-step action"]]
    for s,u,t,o,a in scenarios:
        rows.append([s,u,t,o,a])
    table(rows, widths=[2.6, 0.7, 1.2, 0.8, 2.5])

h2("4. Team tagging  -  Nathalie / Fabi / Noa")
rows=[["Team member","Tag","Role","Territory / scope","After-hours"]]
for n,t,r,c,a in TEAM_TAGS:
    rows.append([n,t,r,c,a])
table(rows, widths=[1.0, 1.0, 2.6, 2.2, 1.4])

body("Rationale recap:")
bullet("Fabi: Apteker's outlying-county agent. Anything outside Nathalie's core territory routes to Fabi by default.")
bullet("Noa: Contracts and back-office admin. All paper, vendor, and money topics route to Noa regardless of county.")
bullet("Nathalie: Owner. Handles relationship work, VIP clients, deals, and is the SOLE after-hours responder for P0 emergencies.")

h2("5. How this feeds the classifier")
body("The classifier prompt is templated from the JSON file (taxonomy.json) alongside this doc. "
     "When the taxonomy is edited, taxonomy.json regenerates the system prompt automatically. The "
     "classifier returns urgency + topic + reasoning; the rules engine then maps to an owner.")

h2("6. Open questions to lock in kickoff")
table([
    ["#",  "Question",                                                                "Default"],
    ["T1", "Are 'showing' and 'deal' really separate topics, or one bucket?",          "Separate (different SLAs)"],
    ["T2", "Should 'vip_client' override topic routing always, or only for buyer/seller?","Always"],
    ["T3", "Tenant issue -> Nathalie or Noa for non-emergency?",                       "Nathalie (keeps relationship)"],
    ["T4", "Confidence floor for taxonomy match (0.55 default)",                       "0.55"],
    ["T5", "Hebrew-language detection -> always Nathalie?",                            "Yes"],
], widths=[0.4, 4.6, 2.0])

doc.save(DOCX)
print("Wrote:", DOCX)
print("Wrote:", JSON)
