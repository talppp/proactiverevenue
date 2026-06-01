"""
Rules and Routing Conditions Draft - the decision logic the AI classifier and
post-classifier rules engine will execute. Emits both DOCX (human-readable) and
YAML (machine-readable, drop-in for the rules engine).
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor

OUT = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router\04_Rules_and_Routing"
DOCX = os.path.join(OUT, "Rules_and_Routing_Conditions_Draft.docx")
YAML = os.path.join(OUT, "routing_rules.yaml")

# ---------------------------------------------------------------------------
# Machine-readable rules
# ---------------------------------------------------------------------------
yaml_text = """# AI Issue Router - Routing Rules v1.0
# Apteker Realty / ProActive Revenue / 2026-05-17
# Evaluation order: filters -> urgency -> topic+geo -> owner -> overflow -> delivery
# This file is read by the Python classifier service (Option C hybrid).

team:
  nathalie:
    role: owner
    counties: [Fulton, DeKalb]              # core territory
    topics:  [buyer_lead, seller_lead, showing, deal, escalation, vip_client]
    after_hours_ok: true                    # accepts P0 outside business hours
  fabi:
    role: agent
    counties: [Cobb, Gwinnett, Forsyth, Cherokee]   # all NON-core counties
    topics:  [buyer_lead, seller_lead, showing, tenant_issue, maintenance]
    after_hours_ok: false
  noa:
    role: back_office
    counties: [ANY]                         # county-agnostic
    topics:  [contract, document, vendor, admin, scheduling, billing]
    after_hours_ok: false

filters:
  # Drop before classifier runs (saves LLM tokens).
  - id: spam_obvious
    when: "regex_match(body, '^(http|view this offer|click here)') AND no_known_sender"
    action: archive
  - id: marketing_promo
    when: "sender_domain in ['mailchimp.com','sendgrid.net','constantcontact.com']"
    action: archive
  - id: personal_keywords
    when: "topic == 'personal' AND confidence > 0.8"
    action: archive_personal
  - id: family_alias
    when: "sender_email in family_aliases_list"
    action: archive_personal

urgency_rules:
  # Output of urgency classifier overrides keyword bumps only upward, never down.
  P0_emergency:
    sla_minutes: 15
    keyword_signals: [emergency, flood, fire, gas leak, no water, broken lock, locked out, burst pipe, no heat, no AC, smoke, child stuck]
    sentiment: panic|alarmed
    deliver: [push_now, sms_owner, escalate_if_unread_5min]
  P1_same_day:
    sla_minutes: 240
    keyword_signals: [today, ASAP, urgent, deadline, closing tomorrow, offer expires]
    deliver: [push_now, queue_card]
  P2_24h:
    sla_minutes: 1440
    keyword_signals: [this week, when you can, soon, please reply]
    deliver: [queue_card, batched_digest]
  P3_backlog:
    sla_minutes: 4320       # 3 days
    keyword_signals: [no rush, fyi, info request, future]
    deliver: [batched_digest]

topic_routing:
  - topic: buyer_lead
    primary_owner: nathalie
    fallback: fabi
    notes: "If county is in fabi's territory and Nathalie is busy (>3 active threads), route to Fabi."
  - topic: seller_lead
    primary_owner: nathalie
    fallback: fabi
  - topic: showing
    primary_owner: nathalie
    fallback: fabi
  - topic: tenant_issue
    primary_owner: nathalie
    fallback: noa
    notes: "Maintenance escalations go through Noa for vendor dispatch."
  - topic: maintenance
    primary_owner: noa
    fallback: fabi
    notes: "Noa coordinates vendor; if vendor not available in county, hands to Fabi."
  - topic: contract
    primary_owner: noa
    fallback: nathalie
  - topic: document
    primary_owner: noa
    fallback: nathalie
  - topic: vendor
    primary_owner: noa
    fallback: nathalie
  - topic: admin
    primary_owner: noa
    fallback: nathalie
  - topic: scheduling
    primary_owner: noa
    fallback: nathalie
  - topic: billing
    primary_owner: noa
    fallback: nathalie
  - topic: deal
    primary_owner: nathalie
    fallback: noa
  - topic: vip_client
    primary_owner: nathalie
    fallback: nathalie       # never offload VIP

geo_overrides:
  # Geography overrides default topic ownership when applicable.
  - when: "county in fabi.counties AND topic in [buyer_lead, seller_lead, showing]"
    route_to: fabi
    notes: "Fabi owns properties outside Nathalie's core counties."
  - when: "topic == contract OR topic == document OR topic == admin"
    route_to: noa
    notes: "Geography does not matter for back-office work."

after_hours_policy:
  start: "19:00"
  end:   "08:00"
  timezone: "America/New_York"
  rules:
    - "P0 -> push to whoever has after_hours_ok=true (Nathalie)."
    - "P1 -> queue, deliver next business morning at 07:30."
    - "P2/P3 -> queue, batch into 09:00 digest."

vip_client_list:
  # Auto-escalate any inbound from these contacts to Nathalie regardless of topic.
  source: "google_contacts label:VIP"
  fallback_to_nathalie: true

deduplication:
  window_seconds: 60
  key: "channel + sender + sha1(body[:200])"

confidence_floor:
  classifier_min_confidence: 0.55
  action_below_floor: "route_to_nathalie_with_unsure_flag"

escalation_paths:
  - condition: "urgency == P0 AND no_human_response_in_minutes >= 5"
    action:    "ring owner phone via Twilio voice prompt"
  - condition: "urgency == P1 AND no_human_response_in_minutes >= 60"
    action:    "second push + cc fallback owner"
"""

with open(YAML, "w", encoding="utf-8") as f:
    f.write(yaml_text)

# ---------------------------------------------------------------------------
# Human-readable DOCX
# ---------------------------------------------------------------------------
doc = Document()
for s in doc.sections:
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)

def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(18)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h2(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(13)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def body(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size=Pt(10)
def bullet(t):
    p = doc.add_paragraph(t, style="List Bullet")
    for r in p.runs: r.font.size=Pt(10)
def table(rows, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0])); t.style="Light Grid Accent 1"
    for i,row in enumerate(rows):
        cells = t.add_row().cells
        for j,val in enumerate(row):
            cells[j].text=str(val)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size=Pt(9)
                    if i==0: r.bold=True
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths):
                row.cells[j].width=Inches(w)

h1("Rules and Routing Conditions  -  Draft v1")
body("Apteker Realty  /  AI Issue Router  /  Phase I  /  Rev 1.0  /  2026-05-17")
body("This document is the human-readable companion to routing_rules.yaml (machine-readable). "
     "Whenever a rule changes, BOTH files must change. The YAML is the source of truth at runtime.")

h2("Evaluation order (pipeline)")
table([
    ["#", "Stage",                          "Purpose"],
    ["1", "Pre-classifier filters",         "Drop obvious spam / personal mail before LLM call"],
    ["2", "Urgency classification",         "Assign P0 / P1 / P2 / P3 with reasoning"],
    ["3", "Topic + Geo classification",     "Tag topic + county + property/listing-id"],
    ["4", "Topic routing table",            "Map topic -> primary owner"],
    ["5", "Geo override",                   "If county is outside Nathalie's core, hand off to Fabi (for lead/showing topics)"],
    ["6", "VIP list check",                 "Force-route to Nathalie if sender is in VIP contacts"],
    ["7", "After-hours policy",             "Hold non-P0 until 07:30 next business morning"],
    ["8", "Confidence floor",               "If classifier confidence < 0.55, default to Nathalie with 'unsure' flag"],
    ["9", "Delivery channel selection",     "Per-urgency delivery (push vs. queue vs. digest)"],
    ["10","Escalation timers",              "If unread / unanswered, escalate per SLA timers"],
], widths=[0.35, 2.4, 4.0])

h2("Team-to-topic primary routing")
table([
    ["Topic",          "Primary owner", "Fallback", "Rationale"],
    ["buyer_lead",      "Nathalie",  "Fabi (if non-core county)", "Nathalie owns lead intake; Fabi for outlying counties"],
    ["seller_lead",     "Nathalie",  "Fabi",                       "Same as above"],
    ["showing",         "Nathalie",  "Fabi",                       "Same as above"],
    ["tenant_issue",    "Nathalie",  "Noa",                        "Nathalie keeps tenant relationship; Noa handles paperwork"],
    ["maintenance",     "Noa",       "Fabi",                       "Vendor coordination is back-office work"],
    ["contract",        "Noa",       "Nathalie",                   "Back-office owns contract review/exec"],
    ["document",        "Noa",       "Nathalie",                   "Same as above"],
    ["vendor",          "Noa",       "Nathalie",                   "Vendor list lives with Noa"],
    ["admin",           "Noa",       "Nathalie",                   "Generic ops"],
    ["scheduling",      "Noa",       "Nathalie",                   "Calendar/showings logistics"],
    ["billing",         "Noa",       "Nathalie",                   "Invoices, deposits, expenses"],
    ["deal",            "Nathalie",  "Noa",                        "Active deals stay with Nathalie"],
    ["vip_client",      "Nathalie",  "Nathalie",                   "Never offload"],
], widths=[1.1, 1.0, 1.8, 3.0])

h2("Geo override - county rules")
body("Nathalie's core counties (defined as her primary territory): Fulton, DeKalb.")
body("Fabi's territory: Cobb, Gwinnett, Forsyth, Cherokee (the 'other counties').")
body("Noa is county-agnostic - back-office tasks route to Noa regardless of where the property sits.")

h2("Urgency tiers and SLA")
table([
    ["Tier", "Name",         "SLA",     "Trigger keywords (examples)",                                       "Delivery"],
    ["P0",   "Emergency",     "15 min",  "emergency, flood, fire, gas leak, no water, locked out, no heat",   "Push + SMS + 5-min unread escalation"],
    ["P1",   "Same-day",      "4 hr",    "today, ASAP, urgent, deadline, closing tomorrow",                  "Push + queue card"],
    ["P2",   "24-hour",       "24 hr",   "this week, when you can, soon",                                    "Queue card + batched digest"],
    ["P3",   "Backlog",       "72 hr",   "no rush, fyi, future",                                             "Batched digest only"],
], widths=[0.45, 1.0, 0.7, 2.7, 2.1])

h2("Edge cases handled in v1")
bullet("Dedupe: same sender + same first 200 chars within 60s -> treat as one message.")
bullet("Confidence floor: classifier confidence < 0.55 -> route to Nathalie tagged 'unsure', she can re-tag.")
bullet("After-hours: 19:00-08:00 ET -> only P0 wakes Nathalie; P1+ batched to 07:30 morning push.")
bullet("VIP override: any sender on the Google Contacts 'VIP' label bypasses topic routing.")
bullet("Multi-language: detect first; if Hebrew, prefer Nathalie regardless of topic (cultural/language reason).")

h2("Edge cases punted to Phase 4 / Phase 5")
bullet("Long thread summarization (Phase 4 - Communication Digest).")
bullet("Auto-learning of routing mistakes (Phase 5 - Memory & Dashboard).")
bullet("Group WhatsApp threads (Phase 4 only - too noisy for Phase I).")

h2("Open decisions to lock in kickoff session")
table([
    ["#",  "Decision",                                                              "Default"],
    ["R1", "Confirm Nathalie's core counties (Fulton + DeKalb correct?)",            "Yes - confirm"],
    ["R2", "Confirm Fabi's counties (Cobb/Gwinnett/Forsyth/Cherokee correct?)",      "Yes - confirm"],
    ["R3", "Is Noa truly county-agnostic, or does she stay with Fulton properties?", "County-agnostic"],
    ["R4", "VIP list source: Google Contacts label or separate spreadsheet?",        "Google Contacts 'VIP'"],
    ["R5", "After-hours window times",                                                "19:00-08:00 ET"],
    ["R6", "P0 escalation: ring phone or just SMS?",                                  "SMS first; ring after 5 min unread"],
    ["R7", "Confidence floor: 0.55 or higher?",                                       "0.55"],
], widths=[0.4, 4.5, 1.8])

doc.save(DOCX)
print("Wrote:", DOCX)
print("Wrote:", YAML)
