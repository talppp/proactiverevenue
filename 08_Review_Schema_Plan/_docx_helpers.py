"""
Shared DOCX helpers that produce accessibility-clean output:
 - Real Heading 1/2/3 styles (so the Word outline + screen reader nav work)
 - First row of every table flagged as a table header (w:tblHeader)
 - Optional alt text + title on images
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BLUE = RGBColor(0x1F, 0x4E, 0x79)

def new_doc(margins=(0.7, 0.7, 0.6, 0.6)):
    d = Document()
    for s in d.sections:
        s.left_margin   = Inches(margins[0])
        s.right_margin  = Inches(margins[1])
        s.top_margin    = Inches(margins[2])
        s.bottom_margin = Inches(margins[3])
    return d

def title(d, text):
    p = d.add_heading(text, level=0)
    for r in p.runs:
        r.font.color.rgb = BLUE
    return p

def h1(d, text):
    p = d.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BLUE
    return p

def h2(d, text):
    p = d.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE
    return p

def h3(d, text):
    p = d.add_heading(text, level=3)
    return p

def body(d, text):
    p = d.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10)
    return p

def bullet(d, text):
    p = d.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(10)
    return p

def numbered(d, text):
    p = d.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(10)
    return p

def code(d, text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    return p

def _mark_table_header(tbl):
    """Set <w:tblHeader/> on the first row so AT announces headers."""
    if not tbl.rows: return
    tr = tbl.rows[0]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    if trPr.find(qn('w:tblHeader')) is None:
        trPr.append(OxmlElement('w:tblHeader'))

def table(d, rows, widths=None, style="Light Grid Accent 1"):
    """First row is header, marked semantically."""
    t = d.add_table(rows=0, cols=len(rows[0]))
    t.style = style
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if i == 0:
                        r.bold = True
    if widths:
        for r in t.rows:
            for j, w in enumerate(widths):
                r.cells[j].width = Inches(w)
    _mark_table_header(t)
    return t

def picture(d, path, width_in, alt_text, alt_title=None):
    """Insert an image with an accessible alt description + title."""
    pic = d.add_picture(path, width=Inches(width_in))
    # Walk back to the <wp:docPr> element and set descr + title.
    p = pic._inline.docPr
    p.set('descr', alt_text)
    if alt_title:
        p.set('title', alt_title)
    return pic
