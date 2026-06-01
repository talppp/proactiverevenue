"""
Quick accessibility inspection of the two deliverables.
Reports: heading structure, table count, image alt text presence, color usage.
"""
from docx import Document
from docx.oxml.ns import qn
import os, re

ROOT = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router"
DOCS = [
    os.path.join(ROOT, r"01_Diagrams\AI_Issue_Router_FlowChart_CAD.docx"),
    os.path.join(ROOT, r"05_Workflow_Mapping\Workflow_Mapping_Document.docx"),
]

def luminance(hexc):
    h = hexc.lstrip("#")
    r,g,b = [int(h[i:i+2],16)/255.0 for i in (0,2,4)]
    def f(c): return c/12.92 if c<=0.03928 else ((c+0.055)/1.055)**2.4
    R,G,B = f(r),f(g),f(b)
    return 0.2126*R + 0.7152*G + 0.0722*B

def contrast(fg, bg):
    L1, L2 = luminance(fg), luminance(bg)
    if L1 < L2: L1, L2 = L2, L1
    return (L1 + 0.05) / (L2 + 0.05)

# Color check
pairs = [
    ("#0A0A0A", "#FFFFFF", "Body black on white"),
    ("#1F4E79", "#FFFFFF", "Blueprint heading on white"),
    ("#C0392B", "#FFFFFF", "Highlight red on white"),
    ("#444444", "#FFFFFF", "Sub-label gray on white"),
    ("#555555", "#FFFFFF", "Title-block label on white"),
    ("#0A0A0A", "#F4F7FA", "Body on lane fill"),
    ("#0A0A0A", "#FFFBEA", "Body on call-out cream"),
    ("#D7DEE3", "#FFFFFF", "Grid line (non-info) on white"),
]
print("=== Color contrast ===")
for fg, bg, label in pairs:
    print(f"  {label:42} {fg} on {bg}: {contrast(fg,bg):.2f}:1")

for p in DOCS:
    print(f"\n=== {os.path.basename(p)} ===")
    d = Document(p)
    # Heading counts
    levels = {}
    for para in d.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            levels[para.style.name] = levels.get(para.style.name,0)+1
    print(f"  Heading distribution: {levels}")
    # Table count
    print(f"  Tables: {len(d.tables)}")
    # Header row marked?
    hdr_marked = 0
    for t in d.tables:
        first = t.rows[0]._tr
        if first.find(qn('w:trPr')) is not None:
            tbl_hdr = first.find(qn('w:trPr')).find(qn('w:tblHeader'))
            if tbl_hdr is not None: hdr_marked += 1
    print(f"  Tables with w:tblHeader on first row: {hdr_marked}/{len(d.tables)}")
    # Images + alt text
    body_xml = d.element.xml
    inlines = re.findall(r'<wp:inline[\s\S]*?</wp:inline>', body_xml) + \
              re.findall(r'<wp:anchor[\s\S]*?</wp:anchor>', body_xml)
    print(f"  Drawing elements: {len(inlines)}")
    for i, inl in enumerate(inlines, 1):
        descr = re.search(r'descr="([^"]*)"', inl)
        title = re.search(r'title="([^"]*)"', inl)
        nv = re.search(r'<wp:docPr [^>]+name="([^"]*)"', inl)
        print(f"    Image {i}: descr={descr.group(1) if descr else '<MISSING>'} | title={title.group(1) if title else '<missing>'} | name={nv.group(1) if nv else '?'}")
