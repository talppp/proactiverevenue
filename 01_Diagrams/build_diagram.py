"""
CAD-style flow chart for the AI Issue Router (Phase I).
Renders a technical-drawing-style diagram (grid, monospace labels, title block,
revision tag, scale bar) and exports PNG + PDF, then embeds into DOCX.
"""

import os
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, Rectangle, FancyBboxPatch, Circle
from matplotlib.lines import Line2D
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router\01_Diagrams"
PNG_PATH = os.path.join(OUT_DIR, "AI_Issue_Router_FlowChart_CAD.png")
PDF_PATH = os.path.join(OUT_DIR, "AI_Issue_Router_FlowChart_CAD.pdf")
DOCX_PATH = os.path.join(OUT_DIR, "AI_Issue_Router_FlowChart_CAD.docx")

# ---------------------------------------------------------------------------
# Style
# ---------------------------------------------------------------------------
plt.rcParams["font.family"] = "DejaVu Sans Mono"
plt.rcParams["font.size"] = 8

INK = "#0A0A0A"
GRID = "#D7DEE3"
BLUEPRINT = "#1F4E79"
LANE_FILL = "#F4F7FA"
CALL_FILL = "#FFFBEA"
HIGHLIGHT = "#C0392B"

fig, ax = plt.subplots(figsize=(20, 13.5))
ax.set_xlim(0, 200)
ax.set_ylim(0, 135)
ax.set_aspect("equal")
ax.axis("off")

# Drawing sheet border (outer + inner)
ax.add_patch(Rectangle((1.5, 1.5), 197, 132, fill=False, lw=2.0, edgecolor=INK))
ax.add_patch(Rectangle((3.0, 3.0), 194, 129, fill=False, lw=0.8, edgecolor=INK))

# Engineering grid
for x in range(5, 200, 5):
    ax.add_line(Line2D([x, x], [3.0, 132.0], color=GRID, lw=0.3, zorder=0))
for y in range(5, 135, 5):
    ax.add_line(Line2D([3.0, 197.0], [y, y], color=GRID, lw=0.3, zorder=0))

# Edge tick markers (CAD coordinate ruler)
for x in range(10, 200, 10):
    ax.text(x, 131.0, str(x), ha="center", va="center", fontsize=6, color=INK)
    ax.text(x, 4.0, str(x), ha="center", va="center", fontsize=6, color=INK)
for y in range(10, 130, 10):
    ax.text(4.0, y, str(y), ha="center", va="center", fontsize=6, color=INK)
    ax.text(196.0, y, str(y), ha="center", va="center", fontsize=6, color=INK)


def draw_box(x, y, w, h, label, sub=None, fill="white", edge=INK, lw=1.2, fontsize=9, italic=False):
    ax.add_patch(Rectangle((x, y), w, h, facecolor=fill, edgecolor=edge, lw=lw, zorder=2))
    style = "italic" if italic else "normal"
    if sub:
        ax.text(x + w / 2, y + h * 0.62, label, ha="center", va="center",
                fontsize=fontsize, fontweight="bold", color=INK, style=style)
        ax.text(x + w / 2, y + h * 0.28, sub, ha="center", va="center",
                fontsize=fontsize - 2, color="#444")
    else:
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=fontsize, fontweight="bold", color=INK, style=style)


def draw_decision(cx, cy, w, h, label):
    # Diamond
    pts = [(cx, cy + h / 2), (cx + w / 2, cy), (cx, cy - h / 2), (cx - w / 2, cy)]
    diamond = plt.Polygon(pts, facecolor="#FFFFFF", edgecolor=INK, lw=1.2, zorder=2)
    ax.add_patch(diamond)
    ax.text(cx, cy, label, ha="center", va="center", fontsize=8, fontweight="bold")


def arrow(x1, y1, x2, y2, label=None, style="-|>", color=INK, lw=1.0, ls="-"):
    a = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle=style, mutation_scale=12,
                        color=color, lw=lw, linestyle=ls, zorder=3)
    ax.add_patch(a)
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.8, label, ha="center", va="bottom",
                fontsize=7, color=color, style="italic")


def lane(y, h, name, color=LANE_FILL):
    ax.add_patch(Rectangle((6, y), 188, h, facecolor=color, edgecolor=INK, lw=0.6, zorder=1))
    ax.text(8, y + h / 2, name, ha="left", va="center", fontsize=9,
            fontweight="bold", color=BLUEPRINT, rotation=0)


def ref(cx, cy, num):
    ax.add_patch(Circle((cx, cy), 1.6, facecolor="white", edgecolor=INK, lw=1.0, zorder=4))
    ax.text(cx, cy, str(num), ha="center", va="center", fontsize=7, fontweight="bold")


# ---------------------------------------------------------------------------
# Title and header strip
# ---------------------------------------------------------------------------
ax.text(100, 127, "AI ISSUE ROUTER  -  PHASE I FLOW DIAGRAM",
        ha="center", va="center", fontsize=16, fontweight="bold", color=BLUEPRINT)
ax.text(100, 123, "APTEKER REALTY  /  PROACTIVE REVENUE  /  REV 1.0  /  2026-05-17",
        ha="center", va="center", fontsize=9, color=INK)

# Horizontal separator
ax.add_line(Line2D([6, 194], [120, 120], color=INK, lw=0.8))

# ---------------------------------------------------------------------------
# Swim lanes
# ---------------------------------------------------------------------------
LANES = [
    (105, 12, "L1  INBOUND CHANNELS"),
    (90,  13, "L2  INGESTION / NORMALIZATION"),
    (72,  16, "L3  AI CLASSIFICATION"),
    (54,  16, "L4  RULES + ROUTING DECISION"),
    (36,  16, "L5  TEAM ASSIGNMENT"),
    (20,  14, "L6  DELIVERY + LOGGING"),
]
for y, h, n in LANES:
    lane(y, h, n)

# ---------------------------------------------------------------------------
# L1 - Inbound channels (5 sources)
# ---------------------------------------------------------------------------
channels = [
    (12,  "APPLE MAIL",    "IMAP /\nMail.app rules"),
    (48,  "WHATSAPP",      "WA Business API\n(Meta-approved)"),
    (84,  "MOBILE SMS",    "Twilio /\nSMS webhook"),
    (120, "INSTAGRAM DM",  "Meta Graph API\n(IG Business)"),
    (156, "GMAIL",         "Gmail API /\nGoogle OAuth"),
]
for i, (x, label, sub) in enumerate(channels, start=1):
    draw_box(x, 108, 32, 7, label, sub, fill="white", fontsize=9)
    ref(x + 32 - 1.8, 108 + 7 - 1.8, i)
    # arrow down to ingestion lane
    arrow(x + 16, 108, x + 16, 103)

# ---------------------------------------------------------------------------
# L2 - Ingestion (single normalizer)
# ---------------------------------------------------------------------------
draw_box(50, 92, 100, 8, "UNIFIED INGESTION BUS",
         "n8n / Make / Zapier  -  webhooks  -  attachment grab  -  language detect  -  dedupe",
         fill=LANE_FILL, fontsize=10)
ref(148, 98, 6)
# fan-in arrows from channels to bus
for x, *_ in channels:
    arrow(x + 16, 103, 100, 100, lw=0.7)
# arrow down to classifier
arrow(100, 92, 100, 88)

# ---------------------------------------------------------------------------
# L3 - Classification (3 parallel boxes feeding to decision)
# ---------------------------------------------------------------------------
draw_box(20, 75, 38, 10, "URGENCY CLASSIFIER",
         "P0 emergency / P1 same-day\nP2 24h / P3 backlog",
         fill="white", fontsize=9)
draw_box(81, 75, 38, 10, "TOPIC CLASSIFIER",
         "Lead / Tenant / Maintenance\nDeal / Vendor / Admin / Spam",
         fill="white", fontsize=9)
draw_box(142, 75, 38, 10, "GEO + PROPERTY TAGGER",
         "County (Fulton/Cobb/...),\nlisting-id, address match",
         fill="white", fontsize=9)
ref(56, 84, 7); ref(117, 84, 8); ref(178, 84, 9)

# Down arrows into decision
arrow(39, 75, 70, 67)
arrow(100, 75, 100, 67)
arrow(161, 75, 130, 67)

# ---------------------------------------------------------------------------
# L4 - Decision diamonds
# ---------------------------------------------------------------------------
draw_decision(50, 62, 22, 10, "URGENCY\n>= P1 ?")
draw_decision(100, 62, 22, 10, "BUSINESS\nRELATED ?")
draw_decision(150, 62, 22, 10, "GEO / TOPIC\nROUTING ?")
ref(61, 67, 10); ref(111, 67, 11); ref(161, 67, 12)

# Yes/No branches
arrow(50, 57, 50, 53, label="YES")
arrow(100, 57, 100, 53, label="YES")
arrow(150, 57, 150, 53, label="YES")

# No paths
arrow(39, 62, 28, 62, label="NO", ls="--")
ax.text(20, 62, "ESCALATE\nQUEUE",
        ha="center", va="center", fontsize=8, fontweight="bold",
        bbox=dict(boxstyle="round,pad=0.3", fc=CALL_FILL, ec=INK))
arrow(89, 62, 80, 58, label="NO", ls="--")
ax.text(76, 56, "ARCHIVE\nLOW PRIO", ha="center", va="center", fontsize=7,
        bbox=dict(boxstyle="round,pad=0.2", fc="#EFEFEF", ec=INK))

# ---------------------------------------------------------------------------
# L5 - Team routing
# ---------------------------------------------------------------------------
draw_box(20, 40, 40, 10, "NATHALIE",
         "Owner / Buyer & Seller\nrelations / VIP escalations", fill="white", fontsize=10)
draw_box(80, 40, 40, 10, "FABI",
         "Properties outside core\ncounty / overflow listings", fill="white", fontsize=10)
draw_box(140, 40, 40, 10, "NOA",
         "Contracts / back office\nadmins / docs / vendor",   fill="white", fontsize=10)
ref(58, 49, 13); ref(118, 49, 14); ref(178, 49, 15)

# Decision -> team mapping arrows
arrow(50, 53, 40, 50)
arrow(100, 53, 100, 50)
arrow(150, 53, 160, 50)

# Cross arrows (overflow)
arrow(60, 45, 80, 45, ls="--")
arrow(120, 45, 140, 45, ls="--")
ax.text(70, 46.5, "overflow", fontsize=6, style="italic")
ax.text(130, 46.5, "overflow", fontsize=6, style="italic")

# ---------------------------------------------------------------------------
# L6 - Delivery / Logging
# ---------------------------------------------------------------------------
draw_box(10, 24, 35, 9, "PUSH\nNOTIFICATION",  "WhatsApp/Slack\nper-owner DM", fill="white", fontsize=8)
draw_box(55, 24, 35, 9, "INBOX\nQUEUE CARDS",   "Trello / Notion\nor CRM card",  fill="white", fontsize=8)
draw_box(100, 24, 35, 9, "AI DRAFT\nREPLY",    "Suggested response\nin-thread",   fill="white", fontsize=8)
draw_box(145, 24, 45, 9, "AUDIT + ANALYTICS\nLOG",
         "Postgres / Sheets  -  feeds Phase 5 Dashboard", fill="white", fontsize=8)
ref(43, 32, 16); ref(88, 32, 17); ref(133, 32, 18); ref(188, 32, 19)

# From team to delivery
for src in [(40, 40), (100, 40), (160, 40)]:
    arrow(src[0], 40, 28, 33, lw=0.6)
    arrow(src[0], 40, 73, 33, lw=0.6)
    arrow(src[0], 40, 118, 33, lw=0.6)
    arrow(src[0], 40, 167, 33, lw=0.6, color=HIGHLIGHT)

# ---------------------------------------------------------------------------
# Title block (CAD-style, bottom-right)
# ---------------------------------------------------------------------------
TB_X, TB_Y, TB_W, TB_H = 130, 6, 64, 12
ax.add_patch(Rectangle((TB_X, TB_Y), TB_W, TB_H, fill=False, lw=1.2, edgecolor=INK))
# Internal grid
for gx in [TB_X + 20, TB_X + 40]:
    ax.add_line(Line2D([gx, gx], [TB_Y, TB_Y + TB_H], color=INK, lw=0.5))
for gy in [TB_Y + 4, TB_Y + 8]:
    ax.add_line(Line2D([TB_X, TB_X + TB_W], [gy, gy], color=INK, lw=0.5))

def tb(col, row, label, value, bold=False):
    cell_w = 20 if col < 2 else 24
    x0 = TB_X + col * 20
    y0 = TB_Y + row * 4
    ax.text(x0 + 0.6, y0 + 2.8, label, ha="left", va="center", fontsize=6, color="#555")
    ax.text(x0 + 0.6, y0 + 1.0, value, ha="left", va="center",
            fontsize=7 + (1 if bold else 0), fontweight="bold" if bold else "normal")

tb(0, 2, "DRAWING NO.",   "APT-AIR-FC-001", bold=True)
tb(1, 2, "REV",           "1.0")
tb(2, 2, "SHEET",         "1 OF 1")
tb(0, 1, "PROJECT",       "AI Issue Router")
tb(1, 1, "PHASE",         "I (Wk 1-2)")
tb(2, 1, "SCALE",         "NOT TO SCALE")
tb(0, 0, "DRAWN",         "ProActive Rev.")
tb(1, 0, "DATE",          "2026-05-17")
tb(2, 0, "APPROVED",      "N. Apteker")

# Legend (bottom-left)
LX, LY = 6, 6
ax.add_patch(Rectangle((LX, LY), 64, 12, fill=False, lw=1.0, edgecolor=INK))
ax.text(LX + 1, LY + 10.5, "LEGEND", fontsize=8, fontweight="bold")
ax.add_patch(Rectangle((LX + 1, LY + 7), 4, 2.5, facecolor="white", edgecolor=INK))
ax.text(LX + 6.5, LY + 8.2, "Process step", fontsize=7)
pts = [(LX + 23, LY + 9.5), (LX + 25, LY + 7.8), (LX + 23, LY + 6.1), (LX + 21, LY + 7.8)]
ax.add_patch(plt.Polygon(pts, fill=False, edgecolor=INK, lw=1.0))
ax.text(LX + 27, LY + 7.8, "Decision", fontsize=7)
ax.add_patch(Circle((LX + 42, LY + 7.8), 1.4, facecolor="white", edgecolor=INK))
ax.text(LX + 42, LY + 7.8, "#", fontsize=6, ha="center", va="center", fontweight="bold")
ax.text(LX + 45, LY + 7.8, "Reference tag", fontsize=7)
ax.add_line(Line2D([LX + 1, LX + 6], [LY + 4.0, LY + 4.0], color=INK, lw=1.0))
ax.text(LX + 7, LY + 4, "Synchronous flow", fontsize=7)
ax.add_line(Line2D([LX + 25, LX + 30], [LY + 4.0, LY + 4.0], color=INK, lw=1.0, ls="--"))
ax.text(LX + 31, LY + 4, "Conditional / fallback", fontsize=7)
ax.add_patch(Rectangle((LX + 1, LY + 1.0), 4, 1.8, facecolor=CALL_FILL, edgecolor=INK))
ax.text(LX + 6.5, LY + 1.8, "Escalation / SLA breach", fontsize=7)

# Scale bar (bottom center)
ax.add_line(Line2D([78, 122], [4.6, 4.6], color=INK, lw=1.2))
for tick in [78, 89, 100, 111, 122]:
    ax.add_line(Line2D([tick, tick], [4.4, 4.8], color=INK, lw=1.0))
ax.text(78, 3.0, "0", fontsize=6, ha="center")
ax.text(100, 3.0, "1 STAGE", fontsize=6, ha="center")
ax.text(122, 3.0, "2 STAGES", fontsize=6, ha="center")

plt.tight_layout()
plt.savefig(PNG_PATH, dpi=220, bbox_inches="tight", facecolor="white")
plt.savefig(PDF_PATH, bbox_inches="tight", facecolor="white")
plt.close(fig)

# ---------------------------------------------------------------------------
# Build a companion DOCX wrapping the diagram with a written legend table
# ---------------------------------------------------------------------------
doc = Document()
section = doc.sections[0]
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI Issue Router  -  Phase I Flow Diagram (CAD-style)")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = meta.add_run("Apteker Realty  -  ProActive Revenue  -  Drawing APT-AIR-FC-001  -  Rev 1.0  -  2026-05-17")
m.font.size = Pt(9)

doc.add_picture(PNG_PATH, width=Inches(7.2))

doc.add_heading("Reference Tags (numbered callouts on drawing)", level=2)

callouts = [
    ("1",  "Apple Mail",          "IMAP push or Mail.app rules export inbound mail for Nathalie's iCloud account."),
    ("2",  "WhatsApp Business",   "Meta-approved BSP webhook (e.g. Twilio, 360dialog) - REQUIRES PROVISIONING."),
    ("3",  "Mobile SMS",          "Twilio long-code / short-code webhook on Nathalie's business mobile number."),
    ("4",  "Instagram DM",        "Meta Graph API on connected IG Business account; inherits Meta App permissions."),
    ("5",  "Gmail",               "Gmail API + Google OAuth on Apteker Realty workspace."),
    ("6",  "Unified Ingestion Bus","Normalizes every payload into one schema {channel, sender, body, attachments, ts, lang}."),
    ("7",  "Urgency Classifier",  "LLM prompt that returns P0-P3 urgency + reasoning."),
    ("8",  "Topic Classifier",    "LLM prompt returning topic + confidence (taxonomy in deliverable 06)."),
    ("9",  "Geo + Property Tagger","Regex + LLM extract: county, listing-id, MLS ref, street address."),
    ("10", "Urgency Gate",        "If urgency >= P1 -> push real-time channel; else queue."),
    ("11", "Business Filter",     "Spam / personal mail diverted to a low-priority archive."),
    ("12", "Geo/Topic Router",    "Selects owner based on county + topic rule table."),
    ("13", "Nathalie",            "Default owner; VIP clients, deals, escalations."),
    ("14", "Fabi",                "Properties OUTSIDE Nathalie's core county (overflow)."),
    ("15", "Noa",                 "Contracts, back-office admin, vendor coordination."),
    ("16", "Push Notification",   "Owner's WhatsApp/Slack/email per delivery preference."),
    ("17", "Inbox Queue",         "CRM card or Trello/Notion task created with full thread context."),
    ("18", "AI Draft Reply",      "Suggested response surfaced inside the original thread."),
    ("19", "Audit + Analytics Log","Postgres/Sheets log feeding Phase 5 Memory & Dashboard."),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "#"
hdr[1].text = "Reference"
hdr[2].text = "Notes"
for tag, ref_name, note in callouts:
    row = tbl.add_row().cells
    row[0].text = tag
    row[1].text = ref_name
    row[2].text = note

doc.add_heading("How to read the swim lanes", level=2)
doc.add_paragraph(
    "L1 INBOUND CHANNELS - the four channels in scope for Phase I "
    "(Apple Mail, WhatsApp, SMS, Instagram) plus Gmail carried over from the original SOW."
)
doc.add_paragraph(
    "L2 INGESTION / NORMALIZATION - one bus, one schema. Every channel's "
    "payload is converted into the same canonical record before AI ever sees it."
)
doc.add_paragraph(
    "L3 AI CLASSIFICATION - three independent passes (urgency, topic, geo). "
    "Runs in parallel for latency."
)
doc.add_paragraph(
    "L4 RULES + ROUTING DECISION - three diamonds that translate AI labels "
    "into a routing decision. All three must pass for a normal route; any 'NO' branches "
    "into an escalation queue or archive."
)
doc.add_paragraph(
    "L5 TEAM ASSIGNMENT - Nathalie / Fabi / Noa. Dashed arrows show overflow "
    "rules when a primary owner is offline or at capacity."
)
doc.add_paragraph(
    "L6 DELIVERY + LOGGING - notification + queue card + AI draft + analytics log. "
    "The analytics log (#19) is the seed of the Phase 5 dashboard."
)

doc.save(DOCX_PATH)
print("Wrote:")
print(" -", PNG_PATH)
print(" -", PDF_PATH)
print(" -", DOCX_PATH)
