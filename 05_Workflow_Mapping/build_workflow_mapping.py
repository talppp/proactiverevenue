"""
Workflow Mapping document - documents the CURRENT (pre-AI) workflow, the FUTURE
(post-AI) workflow, where each Phase I capability slots in, and the data flowing
between each step.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor

OUT = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router\05_Workflow_Mapping"
PATH = os.path.join(OUT, "Workflow_Mapping_Document.docx")

doc = Document()
for s in doc.sections:
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)

def h1(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(18)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h2(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(13)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h3(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(11)
def body(t):
    p=doc.add_paragraph(t)
    for r in p.runs: r.font.size=Pt(10)
def bullet(t):
    p=doc.add_paragraph(t, style="List Bullet")
    for r in p.runs: r.font.size=Pt(10)
def table(rows, widths=None):
    t=doc.add_table(rows=0,cols=len(rows[0])); t.style="Light Grid Accent 1"
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row):
            cells[j].text=str(val)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size=Pt(9)
                    if i==0: r.bold=True
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Inches(w)

h1("Workflow Mapping Document")
body("Apteker Realty  /  AI Issue Router  /  Phase I  /  Rev 1.0  /  2026-05-17")

h2("1. Current state - the manual triage workflow")
body("Before the AI Issue Router, every inbound message hits Nathalie's personal devices first. "
     "She skims, decides, forwards to Fabi or Noa if needed, then re-pings the client. The same "
     "message is often answered twice when threads cross channels (WhatsApp + email).")

table([
    ["#", "Step",                          "Owner",     "Tool",                 "Median time", "Pain"],
    ["1", "Receive message",                "Nathalie",  "Phone (5 apps open)",  "30-60 sec",  "Context switching"],
    ["2", "Read + understand context",      "Nathalie",  "App + memory",         "1-3 min",     "Long threads, voice notes"],
    ["3", "Decide urgency",                 "Nathalie",  "Mental",               "10-30 sec",   "Late-night messages disturb sleep"],
    ["4", "Decide owner (self/Fabi/Noa)",   "Nathalie",  "Mental",               "10-30 sec",   "Bottleneck: Fabi/Noa wait on her"],
    ["5", "Forward / re-explain to owner",  "Nathalie",  "WhatsApp + voice",     "1-2 min",     "Re-typing, summarization"],
    ["6", "Owner responds to client",       "Fabi/Noa",  "Their device",         "2-5 min",     "Back-and-forth on context"],
    ["7", "Status untracked",                "-",         "None",                "-",           "No log; same Q asked again"],
], widths=[0.3, 2.2, 1.0, 1.7, 0.9, 1.6])

body("Quantitative baseline (working assumptions until kickoff):")
bullet("~140 inbound messages/day across all 5 channels.")
bullet("Nathalie spends ~2.5 hours/day on triage alone.")
bullet("~30% of messages are forwarded to Fabi or Noa - that forward takes ~90 seconds each.")
bullet("~10% of messages get a duplicate ('did you see my email?') because nothing tracks status.")

h2("2. Future state - with AI Issue Router")
body("Same channels, same people, but classification, routing, and queueing happen automatically. "
     "Nathalie stops being the human router; she becomes the human EXCEPTION-handler.")

table([
    ["#", "Step",                                "Owner",      "Tool",                 "Median time", "Notes"],
    ["1", "Message enters ingestion bus",         "System",     "n8n",                  "<1 sec",      "Normalized to canonical schema"],
    ["2", "Pre-classifier filter",                "System",     "Rules",                "<1 sec",      "Drops obvious spam/promo"],
    ["3", "AI classify (urgency + topic + geo)",  "Classifier", "Anthropic Claude Haiku","2-4 sec",     "Three parallel calls"],
    ["4", "Rules engine routes",                  "System",     "YAML rules",           "<1 sec",      "Topic + geo + VIP + after-hours"],
    ["5", "Delivery: push / queue / digest",      "System",     "n8n delivery nodes",   "<1 sec",      "Per-urgency policy"],
    ["6", "Owner sees task in their inbox",       "Nat/Fabi/Noa","WhatsApp + CRM card", "instant",     "Includes AI summary + draft reply"],
    ["7", "Owner responds in original channel",   "Owner",       "Native app",          "30 sec-2 min","Draft reply pre-filled"],
    ["8", "Audit log row written",                "System",      "Postgres",            "<1 sec",      "Feeds Phase 5 dashboard"],
], widths=[0.3, 2.6, 1.0, 1.8, 0.9, 1.4])

h2("3. Capability-to-step mapping (which Phase I feature handles which step)")
table([
    ["SOW capability",                                        "Future-state step(s)"],
    ["Auto-classify by urgency, topic, property, geography",  "Step 3"],
    ["Summarize message and recommend next step",             "Step 6 (AI summary inside card)"],
    ["Route to Nathalie / Fabi / Noa by business rules",      "Step 4"],
    ["Works across WhatsApp Business, SMS, Gmail (+Apple Mail, +Instagram)", "Step 1"],
], widths=[4.2, 2.5])

h2("4. Data contract between stages")
h3("Canonical message record (output of Step 1, input to all later steps)")
body("All downstream steps consume one shape; this isolates channel quirks to the ingestion layer.")
table([
    ["Field",            "Type",      "Example",                                     "Notes"],
    ["msg_id",            "uuid",      "8c2f...e7",                                   "Provider id or generated"],
    ["channel",           "enum",      "whatsapp | sms | gmail | apple_mail | instagram", ""],
    ["sender_handle",     "string",    "+14045551212 / john@gmail.com / @ig_handle",  "Normalized"],
    ["sender_known",      "bool",      "true",                                         "Matched against contacts"],
    ["sender_vip",        "bool",      "false",                                        "From VIP list"],
    ["subject",           "string?",   "Re: 1234 Peachtree showing",                   "Email only"],
    ["body",              "string",    "Hi Nathalie, can we move tomorrow's...",      "Plain text, attachments stripped"],
    ["attachments",       "list[url]", "[s3://.../file.pdf]",                         "Persisted out-of-band"],
    ["lang",              "iso639",    "en | he | es",                                "Detected"],
    ["received_at",       "datetime",  "2026-05-17T14:32:01Z",                        "UTC"],
    ["thread_id",         "string?",   "RE-1234-buyer",                               "If part of an existing thread"],
], widths=[1.3, 0.9, 2.5, 2.0])

h3("Classifier output (output of Step 3, input to Step 4)")
table([
    ["Field",                "Type",     "Example",                "Notes"],
    ["urgency",               "enum",     "P0 | P1 | P2 | P3",      ""],
    ["urgency_confidence",    "float",    "0.82",                   "0..1"],
    ["topic",                 "enum",     "buyer_lead | maintenance | contract | ...", "See taxonomy doc"],
    ["topic_confidence",      "float",    "0.74",                   "0..1"],
    ["county",                "string?",  "Fulton",                  "Inferred from address or listing-id"],
    ["listing_id",            "string?",  "MLS-7728312",             "If a listing was mentioned"],
    ["is_business",           "bool",     "true",                    "False = personal / spam"],
    ["sentiment",             "enum",     "neutral | upset | panic | happy", ""],
    ["summary",               "string",   "Client wants to push the showing to Saturday.", "<=160 chars"],
    ["suggested_reply",       "string",   "Hi John, Saturday at 10 works...",           "Drafted; not sent"],
    ["suggested_owner",       "enum",     "nathalie | fabi | noa",   ""],
    ["reasoning",             "string",   "Tenant maintenance issue routed to Noa for vendor dispatch.", "For audit"],
], widths=[1.3, 0.9, 2.5, 2.0])

h3("Routing decision (output of Step 4)")
table([
    ["Field",          "Type",    "Example",                "Notes"],
    ["owner",           "enum",    "nathalie",                ""],
    ["overflow_to",     "enum?",   "fabi",                    "If primary not available"],
    ["delivery_channels","list",   "[push, queue_card]",      ""],
    ["escalate_at",     "datetime?","2026-05-17T14:47:00Z",   "P0 only"],
    ["rule_id",         "string",  "topic.buyer_lead + geo.cobb -> fabi", "For audit / dashboard"],
], widths=[1.3, 0.9, 2.5, 2.0])

h2("5. Where humans intervene")
bullet("Below confidence floor (0.55): all three owners receive a 'help me re-tag' card.")
bullet("Wrong-route correction: Nathalie/Fabi/Noa can re-assign with one tap. The re-assignment is logged and used (Phase 5) to retune the rules.")
bullet("Out-of-scope topics (e.g. legal threats, press inquiries) -> routed to Nathalie with a 'special handling' tag.")

h2("6. Success metrics (measured in Phase 5 dashboard)")
table([
    ["Metric",                                "Baseline (manual)",  "Target (after Phase I)"],
    ["Median triage time per message",         "60-180 sec (Nathalie's eyes)", "<5 sec (auto)"],
    ["Messages routed without Nathalie touch", "0%",                 ">=70%"],
    ["P0 response time",                       "15-90 min",          "<=15 min"],
    ["Duplicate-question rate",                "~10%",               "<3%"],
    ["Routing accuracy (vs human override)",   "n/a",                ">=90% by end of week 2"],
], widths=[3.2, 2.0, 2.0])

h2("7. Integration touchpoints with later phases")
bullet("Phase 2 Vendor Concierge: maintenance topic messages call the Concierge for next-step guidance before the AI draft reply is generated.")
bullet("Phase 3 Buyer Qualification: 'buyer_lead' topic AFTER first call -> triggers the qualification workflow automatically.")
bullet("Phase 4 Communication Digest: consumes the audit log for thread summaries.")
bullet("Phase 5 Memory & Dashboard: the same audit log is the data warehouse for KPIs and anomaly alerts.")

doc.save(PATH)
print("Wrote:", PATH)
