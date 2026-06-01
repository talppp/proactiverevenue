"""
Inbound Channels Discovery Draft - Apple Mail, WhatsApp, Mobile SMS, Instagram, (Gmail).
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router\03_Discovery"
PATH = os.path.join(OUT, "Inbound_Channels_Discovery_Draft.docx")

doc = Document()
for s in doc.sections:
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)

def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h2(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
def h3(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(11)
def body(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(10)
def bullet(t):
    p = doc.add_paragraph(t, style="List Bullet")
    for r in p.runs: r.font.size = Pt(10)
def table(rows, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0])); t.style="Light Grid Accent 1"
    for i,row in enumerate(rows):
        cells = t.add_row().cells
        for j,val in enumerate(row):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if i==0: r.bold=True
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths):
                row.cells[j].width = Inches(w)

# ---------------------------------------------------------------------------
h1("Inbound Channels Discovery Draft")
body("Apteker Realty  /  AI Issue Router  /  Phase I Discovery output  /  Rev 1.0  /  2026-05-17")
body("Purpose: enumerate every inbound channel the AI Issue Router must consume, document "
     "what we need to access it, the access gating risks, the message metadata we can rely on, "
     "data-privacy considerations, and the open questions that block sign-off.")

# ---------------------------------------------------------------------------
h2("Channel summary at a glance")
table([
    ["Channel", "Approx. volume/day", "Access model", "Approval risk", "Phase I scope?"],
    ["Apple Mail (iCloud)",       "30-80 msgs",  "IMAP / app-specific password",   "LOW",    "YES"],
    ["WhatsApp Business",         "60-200 msgs", "Meta-approved BSP webhook",       "HIGH",   "YES"],
    ["Mobile SMS",                "10-40 msgs",  "Twilio long-code webhook",        "LOW",    "YES"],
    ["Instagram DM",              "5-30 msgs",   "Meta Graph API (IG Business)",    "MEDIUM", "YES"],
    ["Gmail (Apteker Realty)",    "40-100 msgs", "Gmail API + Google OAuth",        "LOW",    "YES (carried from SOW)"],
], widths=[1.7, 1.3, 1.7, 1.0, 1.5])
body("Volume estimates are working assumptions until Nathalie confirms during the kickoff session.")

# ---------------------------------------------------------------------------
h2("1. Apple Mail (iCloud)")
h3("Access model")
bullet("Option 1: IMAP from imap.mail.me.com with an app-specific password (iCloud requires 2FA and a generated password - the account password will NOT work).")
bullet("Option 2: Apple Mail on Nathalie's Mac, with a server-side rule that auto-forwards to a dedicated ingestion alias (e.g. inbound+apple@apteker.com). Simpler, but loses headers if forwarded plain.")
h3("What we need from Nathalie")
bullet("iCloud account address she uses for business.")
bullet("App-specific password generated from appleid.apple.com -> Security.")
bullet("Confirmation that she is comfortable granting IMAP read access (read-only is sufficient).")
h3("Risks and edge cases")
bullet("iCloud IMAP rate limits idle connections - heartbeat must be every <29 minutes.")
bullet("Apple does not sign messages with DKIM headers the same way Gmail does - sender verification is weaker; flag iCloud-to-iCloud spoofing scenarios.")
bullet("Mixed personal/business mail on the same account - the classifier must filter out personal mail (see deliverable 06).")
h3("Open questions")
bullet("Should we ingest the WHOLE iCloud inbox or only specific folders (e.g. 'Apteker' label)?")
bullet("Are there shared family iCloud aliases that should be EXCLUDED from ingestion?")

# ---------------------------------------------------------------------------
h2("2. WhatsApp Business API")
h3("Access model")
bullet("Requires a Meta-approved Business Solution Provider (BSP). Recommended: Twilio, 360dialog, or Meta Cloud API direct.")
bullet("Phone number must be migrated from the consumer WhatsApp app to a Business API account. ONCE MIGRATED, the consumer app on the same number stops working.")
bullet("Webhook URL receives every inbound message as JSON; outbound messages cost ~0.5-1 cent per session message.")
h3("What we need from Nathalie")
bullet("The phone number she wants to use as the business inbox. Confirm she is willing to LOSE consumer WhatsApp on this number, OR provision a NEW number.")
bullet("Meta Business Manager admin access (or her granting ProActive Revenue admin).")
bullet("Display name + business category for Meta approval ('Real estate services').")
bullet("Verified business documents (W-9 + DBA filing) - Meta requires legal proof.")
h3("Risks and edge cases")
bullet("Approval timeline: 5-14 days for Meta verification. THIS IS THE LONGEST POLE - kick off Day 1.")
bullet("WhatsApp's 24-hour session window: outbound free-form messages only allowed for 24h after the customer's last message; otherwise must use a pre-approved template.")
bullet("Voice notes and image attachments arrive as URLs valid for ~14 days - we must download and persist them.")
h3("Open questions")
bullet("Are there existing WhatsApp Broadcast lists or Groups Nathalie uses that need to keep working on a different number?")
bullet("Do clients sometimes send a single thread of 15 voice notes? (Affects digest design - Phase 4.)")

# ---------------------------------------------------------------------------
h2("3. Mobile SMS")
h3("Access model")
bullet("Twilio long-code SMS webhook is the standard. A Twilio number receives texts and POSTs to our webhook.")
bullet("Forwarding from Nathalie's existing mobile number to a Twilio number is POSSIBLE via Verizon/AT&T text-to-email gateways but UNRELIABLE - we will recommend a dedicated Twilio number that Nathalie publishes as her business SMS.")
h3("What we need from Nathalie")
bullet("Decision: dedicated Twilio number (recommended) vs. forwarding her personal mobile.")
bullet("A Twilio account (created by ProActive on her behalf if she authorizes).")
bullet("Approval for monthly Twilio cost (~$1/mo for the number + $0.0079 per inbound SMS).")
h3("Risks and edge cases")
bullet("US 10DLC registration is required for ANY business SMS to/from US carriers since 2023. ~$4 one-time + $2/mo brand fee. Approval takes 1-7 days.")
bullet("MMS (photos in SMS) costs ~3x text and may arrive split across multiple messages.")
bullet("STOP / HELP keywords are auto-handled by Twilio but reduce the message count we see.")
h3("Open questions")
bullet("Do existing buyers/sellers expect to text her personal mobile? Plan for a 30-day overlap with a 'we have moved business texts to X' auto-reply.")

# ---------------------------------------------------------------------------
h2("4. Instagram DM (IG Business account)")
h3("Access model")
bullet("Requires an Instagram BUSINESS account (not Personal or Creator) linked to a Facebook Page.")
bullet("Access via Meta Graph API messaging endpoints. Same Meta App Review process as WhatsApp.")
bullet("Webhook delivers inbound DMs as Messenger-style events.")
h3("What we need from Nathalie")
bullet("IG handle and confirmation it is a Business account (or willingness to convert).")
bullet("The linked Facebook Page name and Page admin access.")
bullet("Permissions to be granted to the Meta App: instagram_basic, instagram_manage_messages, pages_messaging.")
h3("Risks and edge cases")
bullet("Meta requires App Review for instagram_manage_messages - 5-15 business days, app-by-app.")
bullet("Reels comments are NOT DMs - they will not flow through this connector. Future enhancement.")
bullet("Story replies arrive as DMs and may reference a media object that expires in 24h - persist the asset immediately.")
h3("Open questions")
bullet("Do buyers / sellers actually reach out via IG today, or is it mostly marketing? If volume is <5/day, deprioritize.")
bullet("Are we OK ignoring comments and reactions in Phase I, focusing only on direct messages?")

# ---------------------------------------------------------------------------
h2("5. Gmail (carried from original SOW)")
h3("Access model")
bullet("Gmail API with OAuth2 - PubSub push notification or polling every 60s.")
bullet("Service account NOT supported for personal Gmail without domain-wide delegation - we will use OAuth on Nathalie's account.")
h3("What we need from Nathalie")
bullet("Consent screen approval for the Gmail API scopes: gmail.readonly + gmail.send (only if AI Draft Reply will be sent automatically; otherwise gmail.readonly suffices for Phase I).")
bullet("Confirmation of the exact Gmail address used for business (and any aliases).")
h3("Risks and edge cases")
bullet("Gmail history pagination tokens expire after 7 days of inactivity - service must reconnect cleanly.")
bullet("Attachments over 25MB are not delivered inline - the API exposes them via a separate fetch.")

# ---------------------------------------------------------------------------
h2("Cross-cutting concerns")
h3("Privacy and data residency")
bullet("All channels carry PII (names, phone numbers, financial detail). Storage encrypted at rest (AES-256) and in transit (TLS 1.2+).")
bullet("LLM calls: redact or hash phone numbers and email addresses BEFORE classifier call when feasible; keep originals in the audit log only.")
bullet("Retention: default 180 days raw; longer for the analytics log (aggregate only).")
h3("Reliability")
bullet("Each channel needs an idempotency key on the webhook (provider-supplied UUID).")
bullet("Provider downtime: each connector should buffer up to 1h locally and replay; if longer, page the on-call.")
h3("Deliverability and rate limits")
bullet("Outbound deliveries must respect each provider's rate limits (e.g. WhatsApp tiered messaging limits, Twilio carrier filtering).")

# ---------------------------------------------------------------------------
h2("Decisions required from Nathalie before build starts")
table([
    ["#",   "Decision",                                                       "Default if no response"],
    ["D1",  "Apple Mail: IMAP credentials or forwarding alias?",              "IMAP with app-specific password"],
    ["D2",  "WhatsApp: keep existing number or provision new business number?","Provision new number"],
    ["D3",  "SMS: dedicated Twilio number vs personal-number forward?",        "Dedicated Twilio number"],
    ["D4",  "Instagram: in-scope for Phase I or defer to Phase 2?",            "In-scope, but defer if Meta review > 14 days"],
    ["D5",  "Gmail: read-only or read+send?",                                  "Read-only (drafts only) for Phase I"],
    ["D6",  "Retention period for raw inbound messages",                       "180 days"],
    ["D7",  "Personal / family alias EXCLUSIONS for each account",             "None (everything classified)"],
], widths=[0.4, 4.0, 2.7])

doc.save(PATH)
print("Wrote:", PATH)
