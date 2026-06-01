"""
Phase I workflow design options - 3 candidate architectures, scored side-by-side.
Emits a DOCX with comparison tables and recommendation.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\USER\Documents\AI_Audit_Meetings_FILES\Natalie -Relator\Phase1_AI_Issue_Router\02_Workflow_Designs"
PATH = os.path.join(OUT, "AI_Issue_Router_Workflow_Design_Options.docx")

doc = Document()
for s in doc.sections:
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def body(text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(10)

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs: r.font.size = Pt(10)

def table(rows, header=True, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if header and i == 0:
                        r.bold = True
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t

# ---------------------------------------------------------------------------
# Cover / framing
# ---------------------------------------------------------------------------
h1("AI Issue Router  -  Workflow Design Options")
body("Apteker Realty  /  ProActive Revenue  /  Phase I  /  Rev 1.0  /  2026-05-17")
body("This document proposes three candidate workflow architectures for the AI Issue Router. "
     "Each option achieves the same business outcome (classify, route, deliver) but trades off "
     "cost, vendor lock-in, ownership, and time-to-value differently. A scored comparison and a "
     "recommendation appear at the end.")

# ---------------------------------------------------------------------------
# Option A - No-code orchestrator (n8n / Make / Zapier)
# ---------------------------------------------------------------------------
h2("Option A  -  No-code orchestrator (n8n / Make / Zapier)")
body("A visual workflow built in a low-code automation platform. Each channel is a trigger node; "
     "the LLM call is one step; routing is a Switch node; delivery is a per-owner sub-flow.")
h3("Architecture")
bullet("Triggers: Gmail node, IMAP node (Apple Mail), WhatsApp Business node, Twilio SMS webhook, IG webhook.")
bullet("Normalize: Set/Function node maps every payload to {channel, sender, body, attachments, ts}.")
bullet("Classify: HTTP Request to Anthropic/OpenAI with two prompts (urgency, topic+geo).")
bullet("Route: Switch node on (urgency, county, topic) -> Nathalie / Fabi / Noa branches.")
bullet("Deliver: WhatsApp Send + Notion/Trello card create + log row in Google Sheets.")
h3("Strengths")
bullet("Fastest to build - 2-3 days for a working pilot.")
bullet("Nathalie's team can SEE the flow and toggle nodes themselves.")
bullet("Minimal hosting cost (n8n cloud ~$20/mo, Make ~$15/mo).")
h3("Weaknesses")
bullet("Logic gets tangled past ~30 nodes; debugging long branches is painful.")
bullet("Per-execution pricing escalates above ~5k messages/month.")
bullet("Vendor lock-in - exporting to another stack is non-trivial.")

# ---------------------------------------------------------------------------
# Option B - Lightweight backend service (Python/Node)
# ---------------------------------------------------------------------------
h2("Option B  -  Lightweight backend service (Python or Node)")
body("A small custom service hosted on Render/Fly.io/Railway with a Postgres database. "
     "Each channel has a webhook endpoint; an internal worker handles classification and routing.")
h3("Architecture")
bullet("Inbound webhooks: /webhook/whatsapp, /webhook/sms, /webhook/instagram, /webhook/gmail, /webhook/applemail.")
bullet("Queue: Postgres-backed job queue (or Redis) for retry + dedupe.")
bullet("Classifier: single function calling Anthropic claude-haiku-4-5 with structured-output JSON.")
bullet("Rules engine: declarative YAML mapping (county -> owner), version-controlled in git.")
bullet("Delivery: same providers as Option A, called via SDK.")
h3("Strengths")
bullet("Predictable cost: ~$25-40/mo hosting + LLM tokens; flat to thousands of messages.")
bullet("Full ownership - code in Apteker's repo; no platform lock-in.")
bullet("Easier to add Phase 4 (Communication Digest) and Phase 5 (Dashboard) on the same DB.")
h3("Weaknesses")
bullet("Slowest to first running prototype - 5-7 build days.")
bullet("Team can't tweak rules without a pull request unless we build an admin UI (out of Phase I scope).")
bullet("Hosting + uptime is now ProActive's monthly support responsibility.")

# ---------------------------------------------------------------------------
# Option C - Hybrid (no-code surface + code-owned core)
# ---------------------------------------------------------------------------
h2("Option C  -  Hybrid: no-code surface, code-owned core (RECOMMENDED)")
body("Use n8n for inbound triggers and outbound delivery (the parts that change often and benefit "
     "from a visual surface), and a small Python classifier service for the AI + rules engine (the parts "
     "that benefit from version control and tests).")
h3("Architecture")
bullet("n8n cloud handles all 5 channel triggers and all 3 delivery channels (WhatsApp, push, CRM card).")
bullet("Between trigger and delivery, n8n calls one HTTPS endpoint on the Python service: POST /route.")
bullet("Python service does: normalize -> classify -> apply rules -> return owner + draft reply + tags.")
bullet("Rules live in YAML in git; LLM prompts live in code; logs live in Postgres.")
bullet("n8n is the team's 'control panel'; the Python service is the brain.")
h3("Strengths")
bullet("Best of both: team gets a visual surface, ProActive owns the testable core.")
bullet("Phase 4 + 5 can extend the same Python service without re-architecture.")
bullet("Adding a sixth channel later means adding one n8n trigger - no code change.")
h3("Weaknesses")
bullet("Two systems to monitor instead of one (mitigated: n8n has built-in error workflows).")
bullet("Slightly higher integration effort upfront (~1.5 extra days vs pure Option A).")

# ---------------------------------------------------------------------------
# Scored comparison
# ---------------------------------------------------------------------------
h2("Scored comparison (1 = poor, 5 = excellent)")
table([
    ["Criterion (weight)",            "A: No-code", "B: Backend service", "C: Hybrid"],
    ["Time to first pilot (15%)",     "5", "2", "4"],
    ["Monthly run cost @ 5k msg (15%)","2", "5", "4"],
    ["Ownership / portability (15%)", "1", "5", "4"],
    ["Team editability (10%)",        "5", "1", "4"],
    ["Extensibility to Phase 4/5 (15%)","2", "5", "5"],
    ["Auditability / logs (10%)",     "3", "5", "4"],
    ["Reliability / SLO control (10%)","2", "5", "4"],
    ["Vendor lock-in risk (10%)",     "1", "5", "4"],
    ["Weighted total",                "2.65", "4.20", "4.10"],
], widths=[2.6, 1.3, 1.6, 1.3])

# ---------------------------------------------------------------------------
# Recommendation
# ---------------------------------------------------------------------------
h2("Recommendation")
body("Adopt OPTION C (Hybrid) for Phase I. The pure-backend Option B scores marginally higher "
     "on paper but loses the team's ability to inspect or toggle the flow themselves - which is a "
     "real cost in a 3-person firm with no internal engineer. The hybrid keeps Nathalie/Fabi/Noa "
     "in the visual surface they understand while ProActive Revenue retains a testable core that "
     "Phases 4 and 5 can build on without rework.")
body("If we hit a hard time constraint (e.g. WhatsApp Business API approval delays push the build "
     "window inside 1 week), drop temporarily to OPTION A and migrate the classifier to Option C "
     "during Phase 2.")

# ---------------------------------------------------------------------------
# Build sequence (mapped to the contracted 6-phase breakdown)
# ---------------------------------------------------------------------------
h2("Phase I build sequence (Option C)")
table([
    ["Day", "Activity", "Owner", "Output"],
    ["1",   "Channel access audit; Meta / Twilio / Google OAuth scopes confirmed", "ProActive + Nathalie", "Connector status sheet"],
    ["1-2", "Discovery sign-off; routing rules locked",                            "Nathalie / Fabi / Noa", "Rules YAML v1"],
    ["2-3", "n8n triggers wired for all 5 channels; normalizer node",              "ProActive",            "Inbound bus live in staging"],
    ["3-5", "Python classifier service + rules engine + prompts",                  "ProActive",            "POST /route returning owner + tags"],
    ["5-6", "Delivery nodes (WhatsApp push, CRM card, AI draft, audit log)",       "ProActive",            "End-to-end staging path"],
    ["6-8", "Test pack: 100 real samples per channel; tune prompts",               "ProActive + Natalie",  "Routing accuracy >=90%"],
    ["8-9", "Go-live on production tokens; shadow run 24h",                        "Both",                 "Shadow run report"],
    ["9-10","Team training + maintenance handoff",                                 "ProActive",            "Owner manual + access list"],
], widths=[0.5, 3.4, 1.5, 1.6])

doc.save(PATH)
print("Wrote:", PATH)
